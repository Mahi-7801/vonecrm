import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import os

def create_issues_doc():
    doc = docx.Document()

    # Color Palette
    NAVY = RGBColor(15, 23, 42)       # #0F172A
    RED = RGBColor(220, 38, 38)      # #DC2626
    DARK_GRAY = RGBColor(51, 65, 85)  # #334155
    GREEN = RGBColor(22, 163, 74)     # #16A34A
    ORANGE = RGBColor(217, 119, 6)    # #D97706
    LIGHT_BG = "F8FAFC"

    # Margins
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = DARK_GRAY

    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("V ONE CRM / MAHI WHATSAPP PLATFORM\n100% COMPLETE PROJECT AUDIT & ISSUES REPORT")
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = NAVY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Comprehensive Technical Audit: System Deficiencies, Critical Bugs, Completion Status & Production Readiness")
    r_sub.font.size = Pt(11.5)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RED

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Helper Functions
    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = NAVY
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = RED
        return p

    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'))

    def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for margin, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{margin}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    # -------------------------------------------------------------
    # SECTION 1: EXECUTIVE VERDICT: IS THE PROJECT 100% COMPLETE?
    # -------------------------------------------------------------
    add_h1("1. Executive Completion Verdict: Is the Project 100% Complete?")
    
    # Highlight Box Table
    box_table = doc.add_table(rows=1, cols=1)
    box_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = box_table.cell(0, 0)
    set_cell_background(c, "FEF2F2")
    set_cell_margins(c, 150, 150, 180, 180)
    
    bp = c.paragraphs[0]
    r_verdict = bp.add_run("VERDICT: NO, THE PROJECT IS NOT 100% COMPLETE (CURRENTLY ~85% COMPLETE)\n")
    r_verdict.font.bold = True
    r_verdict.font.size = Pt(13)
    r_verdict.font.color.rgb = RED

    r_desc = bp.add_run(
        "While V ONE CRM possesses a rich UI layout, comprehensive navigation, full database schema, Meta API integrations, and Razorpay/SMTP hooks, it CANNOT be classified as 100% complete for production deployment. "
        "Critical architectural bottlenecks—such as synchronous HTTP request blocking during bulk sending, hardcoded template parameter mappings, database status ENUM mismatches between Node and PHP, and negative balance vulnerabilities—must be resolved before launching live commercial campaigns."
    )
    r_desc.font.size = Pt(10.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # -------------------------------------------------------------
    # SECTION 2: DETAILED AUDIT FINDINGS BY CATEGORY
    # -------------------------------------------------------------
    add_h1("2. Detailed Audit Findings & Technical Issues")

    issues_list = [
        {
            "num": "Issue #1",
            "title": "Synchronous HTTP Request Loop in Bulk Sending (Causes 504 Gateway Timeout)",
            "severity": "CRITICAL",
            "location": "server/routes/broadcast.js (Lines 232–374)",
            "problem": "Bulk broadcast requests run sequentially inside a synchronous Express HTTP request handler. Sending to 1,000 contacts takes ~3 minutes. Reverse proxies (Cloudflare, Nginx) drop requests after 30-60 seconds, throwing a 504 Timeout error to the user while the loop runs blindly in the background.",
            "impact": "Users get error popups despite sending in progress; leads to duplicate campaign triggers.",
            "solution": "Implement an Asynchronous Job Queue (BullMQ/Redis or setImmediate background worker) returning an immediate 202 Accepted response with a progress tracker."
        },
        {
            "num": "Issue #2",
            "title": "Hardcoded Template Parameter Substitution Bug ('Value 3')",
            "severity": "HIGH",
            "location": "server/routes/broadcast.js (L250-260), server/server.js (L206-216), PHP Worker (L77-84)",
            "problem": "Parameter 1 {{1}} maps to contact name, Parameter 2 {{2}} maps to phone, and Parameter 3+ {{3}} is hardcoded to literally send the string 'Value 3', 'Value 4' to Meta API.",
            "impact": "Custom variables (Order ID, Expiry Date, Custom CSV values) are overwritten with generic text.",
            "solution": "Implement dynamic parameter mapping from frontend UI payload or custom contact attributes."
        },
        {
            "num": "Issue #3",
            "title": "Database Status ENUM Discrepancy Between Node.js & PHP Schedulers",
            "severity": "HIGH",
            "location": "server/routes/broadcast.js (L442), server/server.js (L153), ProcessScheduledBroadcasts.php (L22)",
            "problem": "Node.js schedules and checks broadcasts with status = 'pending'. PHP scheduler checks status = 'scheduled' and updates to 'processing'. Broadcasts scheduled via PHP are completely ignored by Node.js.",
            "impact": "Scheduled messages fail to send depending on which backend API endpoint scheduled them.",
            "solution": "Standardize the database column ENUM values across Node.js and PHP ('pending', 'processing', 'sent', 'cancelled')."
        },
        {
            "num": "Issue #4",
            "title": "Hardcoded Language Code Fallback ('en_US')",
            "severity": "MEDIUM",
            "location": "server/routes/broadcast.js (L15), server/server.js (L194)",
            "problem": "If a template language is not matched, it defaults to 'en_US'. If the template is approved on Meta in Hindi ('hi_IN'), English ('en'), or Spanish ('es_ES'), Meta API rejects it.",
            "impact": "Meta API returns Error #100 / 132001 (Template does not exist in language en_US).",
            "solution": "Dynamically fetch and store approved language codes directly from Meta Graph API response."
        },
        {
            "num": "Issue #5",
            "title": "Unchecked Prepaid Balance Deduction (Negative Wallet Balance Risk)",
            "severity": "MEDIUM",
            "location": "server/routes/broadcast.js (L66-74 & L342)",
            "problem": "Prepaid wallet balance is checked once before starting the broadcast loop (userBalance <= 0). If a user with ₹5 sends 1,000 messages (cost ₹900), balance goes negative (-₹895).",
            "impact": "Prepaid users can exploit system and send unpaid broadcast messages.",
            "solution": "Calculate total estimated broadcast cost upfront and reserve/hold funds before initiating sending."
        },
        {
            "num": "Issue #6",
            "title": "Database Query Exhaustion Per Recipient",
            "severity": "LOW",
            "location": "server/routes/broadcast.js (L298, L321, L342, L346)",
            "problem": "Executes 4-5 separate SQL queries per recipient inside the loop. 10,000 recipients = 40,000 SQL queries.",
            "impact": "Excessive DB connection overhead and potential MySQL connection pool exhaustion.",
            "solution": "Batch SQL inserts into single INSERT INTO messages VALUES (...) statements in chunks of 100."
        },
        {
            "num": "Issue #7",
            "title": "Unverified Meta App / Developer Sandbox Restrictions",
            "severity": "INFO / META",
            "location": "Meta Developer Portal & Graph API v25.0",
            "problem": "Unverified Meta apps or test phone numbers restrict bulk sending to non-whitelisted numbers (Error 131030).",
            "impact": "Bulk messaging fails for unverified business accounts.",
            "solution": "Complete Meta Business Verification and App Review in Meta Developer Dashboard."
        }
    ]

    for item in issues_list:
        add_h2(f"{item['num']}: {item['title']}")
        tbl = doc.add_table(rows=5, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        meta_rows = [
            ("Severity Level", item['severity']),
            ("Code Location", item['location']),
            ("Root Cause", item['problem']),
            ("Business Impact", item['impact']),
            ("Recommended Solution", item['solution'])
        ]

        for idx, (label, val) in enumerate(meta_rows):
            c0 = tbl.cell(idx, 0)
            c1 = tbl.cell(idx, 1)
            set_cell_margins(c0, 60, 60, 100, 100)
            set_cell_margins(c1, 60, 60, 100, 100)

            if idx % 2 == 1:
                set_cell_background(c0, LIGHT_BG)
                set_cell_background(c1, LIGHT_BG)

            r0 = c0.paragraphs[0].add_run(label)
            r0.font.bold = True
            r0.font.size = Pt(9.5)

            r1 = c1.paragraphs[0].add_run(val)
            r1.font.size = Pt(9.5)
            if label == "Severity Level":
                r1.font.bold = True
                if val == "CRITICAL": r1.font.color.rgb = RED
                elif val == "HIGH": r1.font.color.rgb = ORANGE
                elif val == "MEDIUM": r1.font.color.rgb = NAVY

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # -------------------------------------------------------------
    # SECTION 3: SYSTEM COMPLETION CHECKLIST TABLE
    # -------------------------------------------------------------
    add_h1("3. Feature-by-Feature Completion Checklist")
    p = doc.add_paragraph()
    p.add_run("Below is the complete audit breakdown evaluating every major module of V ONE CRM:")

    checklist_data = [
        ("Authentication & Roles", "100% Complete", "JWT Auth, Login, Signup, Admin Login, Private Routes fully functional.", GREEN),
        ("UI Layout & Navigation", "100% Complete", "Responsive Admin & User sidebars, topbars, dark mode, toast notifications.", GREEN),
        ("Contacts Management", "95% Complete", "CSV Import, manual creation, tagging, segmentation. Needs duplicate phone sanitization.", GREEN),
        ("Template Management", "90% Complete", "Meta Graph API sync, creation modal, approval status tracking. Needs multi-language preview.", ORANGE),
        ("Bulk Messaging Engine", "70% Needs Work", "Synchronous HTTP loop causes timeouts. Hardcoded 'Value 3' parameters. Balance leak.", RED),
        ("Scheduled Broadcasts", "75% Needs Work", "Node vs PHP status ENUM mismatch ('pending' vs 'scheduled').", RED),
        ("Interactive Flow Builder", "85% Complete", "Drag-and-drop nodes, welcome flows, AI bot nodes. Needs fallback on unknown input.", ORANGE),
        ("Groq AI Agent Bot", "90% Complete", "Llama-3 model integration, custom system prompts. Needs API rate-limit fallback.", GREEN),
        ("Live 2-Way Chat Inbox", "90% Complete", "Real-time chat, quick replies, media attachments, unread badge counter.", GREEN),
        ("Platform & Meta Billing", "85% Complete", "Razorpay integration, GST invoices, conversation charges. Needs pre-reservation.", ORANGE),
        ("System Logs & Audit", "100% Complete", "Login history, API logs, error traces, audit trails fully captured.", GREEN)
    ]

    check_table = doc.add_table(rows=len(checklist_data) + 1, cols=4)
    check_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Module / Feature", "Status Verdict", "Audit Findings & Remarks", "Status Indicator"]
    for idx, h in enumerate(headers):
        c = check_table.cell(0, idx)
        set_cell_background(c, "0F172A")
        set_cell_margins(c, 100, 100, 120, 120)
        r = c.paragraphs[0].add_run(h)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9.5)

    for row_i, (mod, stat, desc, col) in enumerate(checklist_data, start=1):
        c0 = check_table.cell(row_i, 0)
        c1 = check_table.cell(row_i, 1)
        c2 = check_table.cell(row_i, 2)
        c3 = check_table.cell(row_i, 3)

        set_cell_margins(c0, 60, 60, 100, 100)
        set_cell_margins(c1, 60, 60, 100, 100)
        set_cell_margins(c2, 60, 60, 100, 100)
        set_cell_margins(c3, 60, 60, 100, 100)

        if row_i % 2 == 1:
            set_cell_background(c0, LIGHT_BG)
            set_cell_background(c1, LIGHT_BG)
            set_cell_background(c2, LIGHT_BG)
            set_cell_background(c3, LIGHT_BG)

        r0 = c0.paragraphs[0].add_run(mod)
        r0.font.bold = True
        r0.font.size = Pt(9)

        r1 = c1.paragraphs[0].add_run(stat)
        r1.font.bold = True
        r1.font.size = Pt(9)
        r1.font.color.rgb = col

        r2 = c2.paragraphs[0].add_run(desc)
        r2.font.size = Pt(8.5)

        indicator = "🟢 READY" if col == GREEN else ("🟡 ACTION NEEDED" if col == ORANGE else "🔴 CRITICAL FIX")
        r3 = c3.paragraphs[0].add_run(indicator)
        r3.font.bold = True
        r3.font.size = Pt(8.5)
        r3.font.color.rgb = col

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # -------------------------------------------------------------
    # SECTION 4: STEP-BY-STEP ACTION PLAN FOR 100% PRODUCTION READINESS
    # -------------------------------------------------------------
    add_h1("4. Step-by-Step Action Plan to Reach 100% Production Readiness")
    p = doc.add_paragraph()
    p.add_run("To bring V ONE CRM from 85% to 100% production readiness, implement the following sequential technical refactoring plan:\n")

    steps = [
        ("Step 1: Refactor Bulk Send to Asynchronous Worker", "Move the sending loop in `server/routes/broadcast.js` out of the HTTP request. Return `202 Accepted` immediately and process background sending via `setImmediate` or Redis queue."),
        ("Step 2: Fix Dynamic Parameter Mapping", "Replace the hardcoded `Value ${index+1}` logic in `broadcast.js` and `server.js` with dynamic variable substitution passed from frontend UI payload or custom contact fields."),
        ("Step 3: Standardize Database Status ENUMs", "Update `scheduled_broadcasts.status` queries across Node.js (`server.js`) and PHP (`ProcessScheduledBroadcasts.php`) to use unified status values: `'pending'`, `'processing'`, `'sent'`, `'cancelled'`),"),
        ("Step 4: Implement Upfront Cost Reservation", "Pre-calculate `total_cost = contact_count * rate` before launching broadcast. Deduct or reserve the balance upfront to prevent negative wallet balances."),
        ("Step 5: Complete Meta Business Verification", "Complete Meta Business Manager verification and App Review in Facebook Developer Portal to remove the Sandbox 131030 recipient restriction.")
    ]

    for title, detail in steps:
        p_step = doc.add_paragraph()
        r_st = p_step.add_run(f"• {title}: ")
        r_st.font.bold = True
        r_st.font.color.rgb = NAVY
        p_step.add_run(detail)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # Save document
    output_path = os.path.join(r"c:\Users\ramya\Downloads\mahi", "V_ONE_CRM_Project_Issues_Audit_Report.docx")
    doc.save(output_path)
    print(f"Successfully created Issues Word document at: {output_path}")

if __name__ == "__main__":
    create_issues_doc()
