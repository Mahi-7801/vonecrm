import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import os

def create_team_tickets_doc():
    doc = docx.Document()

    # Color Scheme
    NAVY = RGBColor(15, 23, 42)       # #0F172A
    RED = RGBColor(220, 38, 38)      # #DC2626
    DARK_GRAY = RGBColor(51, 65, 85)  # #334155
    BLUE = RGBColor(37, 99, 235)      # #2563EB
    GREEN = RGBColor(22, 163, 74)     # #16A34A
    ORANGE = RGBColor(217, 119, 6)    # #D97706
    LIGHT_BG = "F8FAFC"
    CODE_BG = "F1F5F9"

    # Margins
    for s in doc.sections:
        s.top_margin = Inches(0.75)
        s.bottom_margin = Inches(0.75)
        s.left_margin = Inches(0.75)
        s.right_margin = Inches(0.75)

    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = DARK_GRAY

    # Document Header
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("V ONE CRM / MAHI WHATSAPP PLATFORM\nTEAM DEVELOPER TASK TICKETS & STEP-BY-STEP SOLUTIONS")
    r_title.font.size = Pt(20)
    r_title.font.bold = True
    r_title.font.color.rgb = NAVY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Actionable Engineering Tickets, Exact File Locations, Refactoring Specifications & Code Fixes for Team Assignment")
    r_sub.font.size = Pt(11)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RED

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Helpers
    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.font.size = Pt(15)
        r.font.bold = True
        r.font.color.rgb = NAVY
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = RED
        return p

    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'))

    def set_cell_margins(cell, top=70, bottom=70, left=100, right=100):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for margin, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{margin}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    def add_code_block(code_text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        c = tbl.cell(0, 0)
        set_cell_background(c, CODE_BG)
        set_cell_margins(c, 80, 80, 120, 120)
        p = c.paragraphs[0]
        r = p.add_run(code_text)
        r.font.name = 'Consolas'
        r.font.size = Pt(9)
        r.font.color.rgb = NAVY
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # SECTION 1: TEAM ASSIGNMENT OVERVIEW
    add_h1("1. Team Assignment Summary & Ticket Dashboard")
    p = doc.add_paragraph()
    p.add_run("This document contains ready-to-assign engineering task tickets for your development team. Each ticket contains the priority, assigned developer role, affected file paths, root cause diagnosis, and the complete step-by-step code implementation solution.")

    tickets_summary = [
        ("TICKET-01", "Backend Developer", "CRITICAL", "Refactor Bulk Send to Async Background Queue", "server/routes/broadcast.js"),
        ("TICKET-02", "Backend / Fullstack Dev", "HIGH", "Fix Hardcoded Parameter 'Value 3' Substitution Bug", "server/routes/broadcast.js, server.js"),
        ("TICKET-03", "Fullstack / DB Dev", "HIGH", "Fix DB Status ENUM Mismatch ('pending' vs 'scheduled')", "broadcast.js, server.js, PHP Worker"),
        ("TICKET-04", "Backend Developer", "MEDIUM", "Dynamic Language Code Extraction (Fix en_US Fallback)", "server/routes/broadcast.js"),
        ("TICKET-05", "Backend Developer", "HIGH", "Pre-calculate Total Cost & Reserve Prepaid Wallet Balance", "server/routes/broadcast.js"),
        ("TICKET-06", "Database Developer", "LOW", "Optimize DB Bulk Message Logging with Chunked Inserts", "server/routes/broadcast.js"),
        ("TICKET-07", "Frontend Developer", "MEDIUM", "Add Live Broadcast Progress Bar & Polling in UI", "client/src/pages/Broadcast.js"),
        ("TICKET-08", "DevOps / System Admin", "INFO", "Complete Meta App Review & Business Manager Verification", "Meta Developer Portal")
    ]

    t_table = doc.add_table(rows=len(tickets_summary) + 1, cols=5)
    t_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Ticket ID", "Assigned Role", "Priority", "Task Title", "Affected File / Component"]
    for idx, h in enumerate(headers):
        c = t_table.cell(0, idx)
        set_cell_background(c, "0F172A")
        set_cell_margins(c, 90, 90, 110, 110)
        r = c.paragraphs[0].add_run(h)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9.5)

    for row_i, (tid, role, prio, title, fpath) in enumerate(tickets_summary, start=1):
        c0 = t_table.cell(row_i, 0)
        c1 = t_table.cell(row_i, 1)
        c2 = t_table.cell(row_i, 2)
        c3 = t_table.cell(row_i, 3)
        c4 = t_table.cell(row_i, 4)

        set_cell_margins(c0, 60, 60, 100, 100)
        set_cell_margins(c1, 60, 60, 100, 100)
        set_cell_margins(c2, 60, 60, 100, 100)
        set_cell_margins(c3, 60, 60, 100, 100)
        set_cell_margins(c4, 60, 60, 100, 100)

        if row_i % 2 == 1:
            for c in [c0, c1, c2, c3, c4]:
                set_cell_background(c, LIGHT_BG)

        r0 = c0.paragraphs[0].add_run(tid)
        r0.font.bold = True
        r0.font.size = Pt(9)

        r1 = c1.paragraphs[0].add_run(role)
        r1.font.size = Pt(9)

        r2 = c2.paragraphs[0].add_run(prio)
        r2.font.bold = True
        r2.font.size = Pt(9)
        if prio == "CRITICAL": r2.font.color.rgb = RED
        elif prio == "HIGH": r2.font.color.rgb = ORANGE
        elif prio == "MEDIUM": r2.font.color.rgb = BLUE
        else: r2.font.color.rgb = DARK_GRAY

        r3 = c3.paragraphs[0].add_run(title)
        r3.font.size = Pt(9)

        r4 = c4.paragraphs[0].add_run(fpath)
        r4.font.size = Pt(8.5)
        r4.font.color.rgb = RED

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # SECTION 2: INDIVIDUAL TICKET SPECIFICATIONS & CODE FIXES
    add_h1("2. Detailed Developer Task Tickets & Solutions")

    # TICKET 1
    add_h2("🎫 TICKET-01: Refactor Bulk Send to Async Background Queue")
    doc.add_paragraph("• Assigned Role: Backend Developer | Priority: CRITICAL | File: server/routes/broadcast.js")
    doc.add_paragraph("Problem Description: The bulk message sending loop runs synchronously inside the HTTP POST /api/broadcast/send handler. Sending to 1,000 contacts takes over 3 minutes, causing Cloudflare/Nginx/browsers to throw HTTP 504 Gateway Timeout errors.")
    doc.add_paragraph("Step-by-Step Solution Code: Replace synchronous loop with setImmediate background execution and return 202 Accepted immediately:")
    
    code_t1 = """// server/routes/broadcast.js - Refactored /send route
router.post('/send', authMiddleware, async (req, res) => {
  try {
    const { template_id, template_name, contact_ids, tag_filter } = req.body;
    
    // 1. Fetch contacts and user balance
    let contacts = await getContactsForBroadcast(req.user.id, contact_ids, tag_filter);
    if (contacts.length === 0) return res.status(400).json({ error: 'No contacts found' });

    // 2. Pre-calculate cost & reserve wallet funds upfront
    const estimatedCost = contacts.length * 0.90; // Default marketing rate
    const [userRows] = await pool.query('SELECT balance, credit_mode FROM users WHERE id = ?', [req.user.id]);
    if (userRows[0]?.credit_mode === 'prepaid' && parseFloat(userRows[0]?.balance || 0) < estimatedCost) {
      return res.status(402).json({ error: `Insufficient wallet balance. Total cost: ₹${estimatedCost}, Balance: ₹${userRows[0].balance}` });
    }

    // 3. Create broadcast job record in DB
    const [jobResult] = await pool.query(
      'INSERT INTO broadcast_jobs (owner_id, template_name, total_contacts, status) VALUES (?, ?, ?, ?)',
      [req.user.id, template_name || 'Template', contacts.length, 'processing']
    );

    // 4. Return IMMEDIATE 202 Accepted response to UI (prevents 504 timeout)
    res.status(202).json({
      message: 'Broadcast job started in background',
      job_id: jobResult.insertId,
      total_contacts: contacts.length,
      status: 'processing'
    });

    // 5. Asynchronous Background Execution
    setImmediate(async () => {
      let sent = 0, failed = 0;
      for (const contact of contacts) {
        try {
          await sendSingleWhatsAppTemplate(req.user.id, contact, template_name, template_id);
          sent++;
        } catch (err) {
          failed++;
        }
        await new Promise(r => setTimeout(r, 20)); // Rate limit pause
      }
      await pool.query('UPDATE broadcast_jobs SET status = "completed", sent_count = ?, failed_count = ? WHERE id = ?',
        [sent, failed, jobResult.insertId]);
    });
  } catch (err) {
    res.status(500).json({ error: err.message });
  }
});"""
    add_code_block(code_t1)

    # TICKET 2
    add_h2("🎫 TICKET-02: Fix Hardcoded Parameter 'Value 3' Substitution Bug")
    doc.add_paragraph("• Assigned Role: Backend / Fullstack Developer | Priority: HIGH | Files: server/routes/broadcast.js, server/server.js")
    doc.add_paragraph("Problem Description: In parameter mapping, index 0 maps to name, index 1 maps to phone, and any index >= 2 is hardcoded to output 'Value 3', 'Value 4', overwriting custom customer variables.")
    doc.add_paragraph("Step-by-Step Solution Code: Update parameter mapping function to accept custom parameter payloads or fallback dynamically to custom contact attributes:")
    
    code_t2 = """// Dynamic Parameter Extraction & Mapping Function
const buildTemplateParameters = (templateParams, contact, customParamValues = {}) => {
  if (!templateParams || templateParams.length === 0) return [];

  return templateParams.map((param, index) => {
    const paramKey = `var_${index + 1}`; // e.g. var_1, var_2, var_3
    let value = '';

    // 1. Priority 1: User payload explicitly provided mapping from UI
    if (customParamValues[paramKey]) {
      value = customParamValues[paramKey];
    }
    // 2. Priority 2: Standard Contact Attributes
    else if (index === 0 && contact.name) {
      value = contact.name;
    } else if (index === 1 && contact.phone) {
      value = contact.phone;
    } else if (contact.company && index === 2) {
      value = contact.company;
    } else if (contact.custom_fields && contact.custom_fields[`var_${index + 1}`]) {
      value = contact.custom_fields[`var_${index + 1}`];
    } 
    // 3. Fallback: Customer Name or blank space (Never hardcode 'Value N')
    else {
      value = contact.name || 'Customer';
    }

    return { type: 'text', text: String(value) };
  });
};"""
    add_code_block(code_t2)

    # TICKET 3
    add_h2("🎫 TICKET-03: Fix DB Status ENUM Mismatch ('pending' vs 'scheduled')")
    doc.add_paragraph("• Assigned Role: Fullstack / DB Developer | Priority: HIGH | Files: broadcast.js, server.js, ProcessScheduledBroadcasts.php")
    doc.add_paragraph("Problem Description: Node.js scheduler checks status = 'pending', whereas PHP worker checks status = 'scheduled'. Scheduled broadcasts created in PHP API are ignored by Node.")
    doc.add_paragraph("Step-by-Step Solution Code: Standardize ENUM in MySQL database and update queries across both Node.js and PHP:")

    code_t3 = """-- MySQL Database Migration
ALTER TABLE scheduled_broadcasts 
MODIFY COLUMN status ENUM('pending', 'processing', 'sent', 'cancelled') DEFAULT 'pending';

-- Node.js server.js Scheduler Query Update
const [due] = await pool.query(
  "SELECT * FROM scheduled_broadcasts WHERE status = 'pending' AND scheduled_at <= NOW()"
);

// PHP ProcessScheduledBroadcasts.php Query Update
$broadcasts = ScheduledBroadcast::where('status', 'pending')
    ->where('scheduled_at', '<=', $now)
    ->get();"""
    add_code_block(code_t3)

    # TICKET 4
    add_h2("🎫 TICKET-04: Dynamic Language Code Extraction (Fix en_US Fallback)")
    doc.add_paragraph("• Assigned Role: Backend Developer | Priority: MEDIUM | File: server/routes/broadcast.js")
    doc.add_paragraph("Problem Description: Hardcoded default to 'en_US' causes Meta API rejection (Error 132001) for templates approved in Hindi (hi_IN), English (en), or Spanish (es_ES).")
    doc.add_paragraph("Step-by-Step Solution Code: Extract approved language directly from Meta response:")

    code_t4 = """// Extract exact language code from Meta API Template object
if (foundMetaTemplate && foundMetaTemplate.language) {
  templateLanguageCode = foundMetaTemplate.language;
} else if (localTemplate && localTemplate.language) {
  templateLanguageCode = localTemplate.language;
} else {
  templateLanguageCode = 'en'; // Meta default standard language code
}

console.log(`Using validated Meta template language: ${templateLanguageCode}`);"""
    add_code_block(code_t4)

    # TICKET 5
    add_h2("🎫 TICKET-05: Pre-calculate Cost & Reserve Wallet Balance")
    doc.add_paragraph("• Assigned Role: Backend Developer | Priority: HIGH | File: server/routes/broadcast.js")
    doc.add_paragraph("Problem Description: Balance check is performed once before loop; user balance can drop into negative values during sending.")
    doc.add_paragraph("Step-by-Step Solution Code: Pre-calculate total cost and reserve balance upfront:")

    code_t5 = """// Pre-calculate total campaign cost
const rate = 0.90; // Marketing rate per message
const totalCampaignCost = contacts.length * rate;

const [userRows] = await pool.query('SELECT balance, credit_mode FROM users WHERE id = ?', [req.user.id]);
const userBalance = parseFloat(userRows[0]?.balance || 0);

if (userRows[0]?.credit_mode === 'prepaid' && userBalance < totalCampaignCost) {
  return res.status(402).json({
    error: `Insufficient wallet balance. Total campaign cost for ${contacts.length} contacts is ₹${totalCampaignCost.toFixed(2)}, but current balance is ₹${userBalance.toFixed(2)}.`
  });
}

// Pre-reserve funds upfront
if (userRows[0]?.credit_mode === 'prepaid') {
  await pool.query('UPDATE users SET balance = balance - ? WHERE id = ?', [totalCampaignCost, req.user.id]);
}"""
    add_code_block(code_t5)

    # TICKET 6
    add_h2("🎫 TICKET-06: Optimize DB Bulk Message Logging with Chunked Inserts")
    doc.add_paragraph("• Assigned Role: Database Developer | Priority: LOW | File: server/routes/broadcast.js")
    doc.add_paragraph("Problem Description: 4-5 single SQL queries per contact inside sending loop overheads MySQL connection pool.")
    doc.add_paragraph("Step-by-Step Solution Code: Use chunked multi-row SQL INSERT queries:")

    code_t6 = """// Chunk array helper
const chunkArray = (arr, size) => Array.from({ length: Math.ceil(arr.length / size) }, (v, i) => arr.slice(i * size, i * size + size));

// Bulk insert message logs in batches of 100
const messageBatches = chunkArray(messageLogRecords, 100);
for (const batch of messageBatches) {
  const values = batch.map(r => [r.owner_id, r.contact_id, 'outbound', r.body, r.template_id, r.wa_message_id, r.status]);
  await pool.query(
    'INSERT INTO messages (owner_id, contact_id, direction, body, template_id, wa_message_id, status) VALUES ?',
    [values]
  );
}"""
    add_code_block(code_t6)

    # TICKET 7
    add_h2("🎫 TICKET-07: Add Live Broadcast Progress Bar & Polling in UI")
    doc.add_paragraph("• Assigned Role: Frontend Developer | Priority: MEDIUM | File: client/src/pages/Broadcast.js")
    doc.add_paragraph("Problem Description: Frontend UI currently waits for long HTTP POST response without progress feedback.")
    doc.add_paragraph("Step-by-Step Solution Code: Add polling for background job status:")

    code_t7 = """// Frontend polling implementation in client/src/pages/Broadcast.js
const handleSend = async () => {
  setSending(true);
  try {
    const res = await api.post('/broadcast/send', payload);
    const jobId = res.data.job_id;
    
    // Poll job status every 2 seconds
    const interval = setInterval(async () => {
      const statusRes = await api.get(`/broadcast/job/${jobId}`);
      setResult(statusRes.data);
      if (statusRes.data.status === 'completed' || statusRes.data.status === 'failed') {
        clearInterval(interval);
        setSending(false);
      }
    }, 2000);
  } catch (err) {
    alert(err.response?.data?.error || 'Broadcast failed');
    setSending(false);
  }
};"""
    add_code_block(code_t7)

    # TICKET 8
    add_h2("🎫 TICKET-08: Complete Meta App Review & Business Manager Verification")
    doc.add_paragraph("• Assigned Role: DevOps / System Administrator | Priority: INFO | Meta Developer Portal")
    doc.add_paragraph("Problem Description: Sandbox mode blocks broadcast messages to non-test phone numbers (Meta Error 131030).")
    doc.add_paragraph("Steps to Complete:")
    doc.add_paragraph("1. Log into Meta Developer Portal (https://developers.facebook.com).\n2. Select WhatsApp App -> Go to Settings -> Basic -> Fill Privacy Policy URL & Terms of Service.\n3. Navigate to Business Manager Settings -> Security Center -> Complete Business Verification.\n4. Submit WhatsApp Business Management & Messaging permission scopes for App Review.")

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # Save document
    output_path = os.path.join(r"c:\Users\ramya\Downloads\mahi", "V_ONE_CRM_Developer_Task_Tickets_And_Solutions.docx")
    doc.save(output_path)
    print(f"Successfully created Team Tickets Word document at: {output_path}")

if __name__ == "__main__":
    create_team_tickets_doc()
