import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import os

def create_document():
    doc = docx.Document()

    # Define color scheme
    NAVY = RGBColor(15, 23, 42)      # #0F172A
    RED = RGBColor(220, 38, 38)     # #DC2626
    DARK_GRAY = RGBColor(51, 65, 85) # #334155
    LIGHT_BG = "F8FAFC"
    BORDER_COLOR = "CBD5E1"

    # Set Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles helper
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = DARK_GRAY

    # Document Header Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("V ONE CRM / MAHI WHATSAPP PLATFORM\nCOMPLETE SYSTEM DOCUMENTATION")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = NAVY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Comprehensive Technical & Functional Guide: Admin Flow, User Flow, Navbar Functions & Required API Keys")
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RED

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Helper function for headings
    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = NAVY
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RED
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = NAVY
        return p

    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'))

    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for margin, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{margin}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    # -------------------------------------------------------------
    # SECTION 1: SYSTEM OVERVIEW & ARCHITECTURE
    # -------------------------------------------------------------
    add_h1("1. System Overview & Technical Architecture")
    p = doc.add_paragraph()
    p.add_run("V ONE CRM (Mahi WhatsApp Marketing & CRM) is a modern, enterprise-grade multi-tenant WhatsApp Business Cloud API marketing, automation, and live CRM platform. It empowers businesses to run mass broadcast campaigns, automate customer service using Groq AI chatbots and interactive drag-and-drop flow builders, manage multi-agent 2-way chat inboxes, and process real-time payments using Meta Cloud API and Razorpay.")
    
    add_h2("Core Technology Stack")
    tech_table = doc.add_table(rows=5, cols=2)
    tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Layer / Component", "Technology & Framework Used"]
    
    for i, h in enumerate(headers):
        cell = tech_table.cell(0, i)
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, 120, 120, 150, 150)
        r = cell.paragraphs[0].add_run(h)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)

    stack_data = [
        ("Frontend Application", "React.js (v18), React Router (v6/v7), Custom CSS Design System, React Icons, React Toastify"),
        ("Backend REST API", "Node.js, Express.js framework, JWT Authentication, Multer file upload, Axios"),
        ("Database Engine", "MySQL Relational Database (whatsapp_crm) storing users, WABA accounts, templates, campaigns, contacts, drip steps, logs"),
        ("Integrations & Services", "Meta WhatsApp Cloud API (v25.0), Meta Embedded Signup SDK, Groq AI (Llama-3), Razorpay Payment Gateway, Nodemailer SMTP")
    ]

    for row_idx, (layer, tech) in enumerate(stack_data, start=1):
        c0 = tech_table.cell(row_idx, 0)
        c1 = tech_table.cell(row_idx, 1)
        set_cell_margins(c0, 100, 100, 150, 150)
        set_cell_margins(c1, 100, 100, 150, 150)
        if row_idx % 2 == 1:
            set_cell_background(c0, LIGHT_BG)
            set_cell_background(c1, LIGHT_BG)
        c0.paragraphs[0].add_run(layer).font.bold = True
        c1.paragraphs[0].add_run(tech)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # -------------------------------------------------------------
    # SECTION 2: ADMIN FLOW & NAVBAR WORK BREAKDOWN
    # -------------------------------------------------------------
    add_h1("2. Admin Flow & Admin Navbar Work Breakdown")
    p = doc.add_paragraph()
    p.add_run("The Admin Flow provides executive and operational oversight across the entire platform. The Super Admin manages platform client accounts, monitors WhatsApp Cloud API WABA account statuses, oversees Meta conversation costs, supervises message broadcast campaigns, configures AI agents, and manages platform subscription billing.")

    add_h2("Admin Access & Route Architecture")
    p = doc.add_paragraph()
    p.add_run("• Admin Login Endpoint: ").font.bold = True
    p.add_run("/admin/login (Authenticates with role = 'admin')\n")
    p.add_run("• Admin Console Root: ").font.bold = True
    p.add_run("/admin (Renders Executive Platform Console)\n")
    p.add_run("• Guarded Layout: ").font.bold = True
    p.add_run("Layout component dynamically renders dark-themed collapsible Admin Navbar with red highlighting (#DC2626).")

    add_h2("Admin Navbar Breakdown & Work Functions")
    p = doc.add_paragraph()
    p.add_run("Below is the detailed work breakdown for each section and each navbar item available in the Admin interface:")

    admin_nav_data = [
        ("📊 Overview", [
            ("Super Dashboard", "/admin", "Displays high-level platform KPIs: total clients, active WhatsApp numbers, messages sent today, total platform revenue, total wallet balances, and live server status."),
            ("Platform Users", "/admin/users?tab=all", "Provides a comprehensive table of all platform users. Admins can search by email/phone/WABA ID, filter by role/status, download full platform CSV reports, suspend abusing accounts, and send manual/automated plan expiry alerts."),
            ("WA Numbers", "/admin/numbers?tab=all", "Lists all connected WhatsApp phone numbers across all client accounts with phone number IDs, verified display names, quality ratings, and status.")
        ]),
        ("👥 User Management", [
            ("All Users", "/admin/users?tab=all", "Main user management view to edit user profiles, assign credit modes (Prepaid vs. Postpaid), and manage user access."),
            ("User Profiles", "/admin/users?tab=profiles", "Inspect individual user usage metrics, connected WABA details, monthly spend, and campaign statistics."),
            ("Roles & Permissions", "/admin/users?tab=roles", "Manage administrative role delegation (Super Admin, Manager, Agent, Client User)."),
            ("Suspended Users", "/admin/users?tab=suspended", "Filter view of suspended user accounts with one-click restore / unsuspend functionality.")
        ]),
        ("📘 Facebook Integration", [
            ("FB Connections", "/admin/facebook?tab=connections", "Monitors live Meta Embedded Signup OAuth connections and Facebook user sessions."),
            ("Page Access Tokens", "/admin/facebook?tab=tokens", "View and refresh Facebook Page Access Tokens associated with connected apps."),
            ("Token Expiry Status", "/admin/facebook?tab=expiry", "Tracks access token expiration timelines and sends background alerts before tokens expire."),
            ("Business Manager", "/admin/facebook?tab=bm", "Manages Meta Business Manager account bindings and organization verification statuses."),
            ("Permissions Review", "/admin/facebook?tab=permissions", "Audits granted Meta Graph API scope permissions (e.g. whatsapp_business_messaging).")
        ]),
        ("💬 WhatsApp Cloud API", [
            ("WABA Accounts", "/admin/whatsapp-api?tab=waba", "Displays all WABA (WhatsApp Business Account) IDs, default system phone number IDs, and permanent system user tokens."),
            ("Quality Ratings", "/admin/whatsapp-api?tab=quality", "Monitors Meta phone number quality scores (GREEN, YELLOW, RED) and messaging volume limit tiers (TIER_250, TIER_1K, TIER_10K, TIER_UNLIMITED)."),
            ("Webhook Status", "/admin/whatsapp-api?tab=webhook", "Live monitor for Meta Webhook endpoint (/api/whatsapp/webhook), verifying signature headers and inbound message processing."),
            ("API Health", "/admin/whatsapp-api?tab=health", "Executes ping tests against Meta Graph API (v25.0) to measure latency and service availability."),
            ("Token Status", "/admin/whatsapp-api?tab=token", "Validates System User Access Tokens and alerts when token refresh is required.")
        ]),
        ("📢 Campaigns & Broadcast", [
            ("All Campaigns", "/admin/messages?tab=all", "Monitors all mass message broadcast campaigns executed by users across the system."),
            ("Delivery Reports", "/admin/messages?tab=reports", "Granular log of message delivery statuses (sent, delivered, read, failed)."),
            ("Campaign Analytics", "/admin/messages?tab=analytics", "System-wide broadcast performance charts, response rates, and delivery success metrics.")
        ]),
        ("📝 Templates", [
            ("All Templates", "/admin/templates?tab=all", "Central repository of all WhatsApp message templates created across all tenant accounts."),
            ("Pending Approval", "/admin/templates?tab=pending", "Queue of message templates submitted to Meta awaiting review."),
            ("Approved", "/admin/templates?tab=approved", "List of active Meta-approved templates ready for broadcast sending."),
            ("Rejected", "/admin/templates?tab=rejected", "List of templates rejected by Meta with detailed rejection reason codes.")
        ]),
        ("🔄 Flow Builder", [
            ("All Flows", "/admin/flows?tab=all", "View and manage visual multi-step interactive chatbot flows created on the platform."),
            ("Published Flows", "/admin/flows?tab=published", "List of active, live interactive flows currently handling incoming customer chats."),
            ("Draft Flows", "/admin/flows?tab=draft", "Work-in-progress flow configurations.")
        ]),
        ("🤖 AI Agents", [
            ("All Agents", "/admin/agents", "Configures Groq-powered AI chatbot agents (Customer Support, Sales Bot, Lead Qualifier), system prompts, model parameters, and auto-reply behavior.")
        ]),
        ("📅 Drip Sequences", [
            ("All Sequences", "/admin/drip-sequences", "Manages automated multi-step time-delayed drip messaging sequences across tenant accounts.")
        ]),
        ("💳 Platform Billing", [
            ("Subscription Plans", "/admin/pricing?tab=plans", "Create and edit platform subscription tiers (Starter, Pro, Enterprise) with pricing and feature limits."),
            ("All Transactions", "/admin/pricing?tab=transactions", "Complete financial ledger of user wallet top-ups, subscription fees, and balance deductions."),
            ("Razorpay Payments", "/admin/pricing?tab=razorpay", "Integration log for Razorpay payment gateway orders, payments, signatures, and webhooks."),
            ("Refunds", "/admin/pricing?tab=refunds", "Process refund requests and manage wallet balance reversals."),
            ("GST Invoices", "/admin/pricing?tab=gst", "Generate tax-compliant GST invoices for user transactions.")
        ]),
        ("💰 Meta Billing", [
            ("Conversation Charges", "/admin/pricing?tab=conversation", "Tracks direct Meta conversation costs broken down by category (Marketing, Utility, Service, Authentication)."),
            ("Marketing Charges", "/admin/pricing?tab=marketing", "Detailed spend analytics for Meta Marketing conversations."),
            ("Utility Charges", "/admin/pricing?tab=utility", "Detailed spend analytics for Meta Utility conversations."),
            ("Auth Charges", "/admin/pricing?tab=auth", "Detailed spend analytics for Meta Authentication OTP conversations."),
            ("Meta Invoices", "/admin/pricing?tab=invoices", "Reconciliation tool matching Meta monthly statements with internal platform message logs.")
        ]),
        ("📋 Contacts & CRM", [
            ("All Contacts", "/admin/contacts?tab=all", "Global database view of all contacts stored across user accounts."),
            ("Segments & Labels", "/admin/contacts?tab=labels", "System-wide contact labels, tags, and audience segmentation."),
            ("Opt-out List", "/admin/contacts?tab=optout", "Centralized list of unsubscribed contacts ('STOP' triggers) ensuring Meta compliance.")
        ]),
        ("📬 Notifications", [
            ("System Alerts", "/admin/notifications-admin?tab=alerts", "High-priority administrative alerts regarding API failures, low balances, or system errors."),
            ("Email Notifications", "/admin/notifications-admin?tab=email", "Logs of transactional email alerts sent via SMTP."),
            ("WA Notifications", "/admin/notifications-admin?tab=whatsapp", "Automated system notification logs sent via WhatsApp.")
        ]),
        ("📜 Activity & Logs", [
            ("Login History", "/admin/logs?tab=login", "Security audit log tracking user & admin login IPs, timestamps, and browser user-agents."),
            ("API Logs", "/admin/logs?tab=api", "Real-time log of REST API calls, endpoints hit, and response status codes."),
            ("Error Logs", "/admin/logs?tab=error", "Exception stack traces and error logs for fast debugging."),
            ("Audit Trail", "/admin/logs?tab=audit", "Historic record of administrative actions, config edits, and balance adjustments.")
        ]),
        ("⚙️ System Settings", [
            ("General Settings", "/admin/settings?tab=general", "Configures application name, branding logo, default timezone, default currency, and contact info."),
            ("API Configuration", "/admin/settings?tab=api", "Global key management for Meta Graph API, Groq AI API, and Razorpay Gateway."),
            ("SMTP Settings", "/admin/settings?tab=smtp", "Configures mail server settings (Host, Port, Username, App Password, From Name)."),
            ("Webhook Settings", "/admin/settings?tab=webhook", "Configures public Webhook URL and Meta Webhook Verification Token."),
            ("Server & Cache", "/admin/settings?tab=server", "Monitors Node.js process uptime, RAM usage, and provides cache flush controls."),
            ("Backup & Restore", "/admin/settings?tab=backup", "Generates full MySQL database backups and snapshot exports."),
            ("Plans Manager", "/admin/plans", "Manage subscription plan boundaries and feature flags.")
        ])
    ]

    for section_title, items in admin_nav_data:
        add_h3(section_title)
        tbl = doc.add_table(rows=len(items) + 1, cols=3)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header
        h_cols = ["Navbar Item Name", "Route Path / Query", "Functionality & Work Done"]
        for idx, htext in enumerate(h_cols):
            c = tbl.cell(0, idx)
            set_cell_background(c, "0F172A")
            set_cell_margins(c, 100, 100, 120, 120)
            r = c.paragraphs[0].add_run(htext)
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(9.5)

        for row_i, (item_name, route, desc) in enumerate(items, start=1):
            c0 = tbl.cell(row_i, 0)
            c1 = tbl.cell(row_i, 1)
            c2 = tbl.cell(row_i, 2)
            set_cell_margins(c0, 80, 80, 100, 100)
            set_cell_margins(c1, 80, 80, 100, 100)
            set_cell_margins(c2, 80, 80, 100, 100)
            if row_i % 2 == 1:
                set_cell_background(c0, LIGHT_BG)
                set_cell_background(c1, LIGHT_BG)
                set_cell_background(c2, LIGHT_BG)

            r0 = c0.paragraphs[0].add_run(item_name)
            r0.font.bold = True
            r0.font.size = Pt(9.5)
            
            r1 = c1.paragraphs[0].add_run(route)
            r1.font.size = Pt(9)
            r1.font.color.rgb = RED
            
            r2 = c2.paragraphs[0].add_run(desc)
            r2.font.size = Pt(9)

        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # -------------------------------------------------------------
    # SECTION 3: USER FLOW & NAVBAR WORK BREAKDOWN
    # -------------------------------------------------------------
    add_h1("3. User Flow & User Navbar Work Breakdown")
    p = doc.add_paragraph()
    p.add_run("The User Flow represents the end-to-end customer journey for business account owners and team agents using V ONE CRM for WhatsApp marketing, CRM, and automation. A typical user logs in, connects their Meta WhatsApp Business Account via Embedded Signup, recharges their wallet, creates Meta-approved templates, imports contacts, launches broadcast campaigns, and handles incoming customer conversations in real-time.")

    add_h2("User Access & Onboarding Journey")
    p = doc.add_paragraph()
    p.add_run("1. Registration & Login: ").font.bold = True
    p.add_run("User signs up at /signup and logs in at /login (JWT token stored in localStorage).\n")
    p.add_run("2. Meta WhatsApp Onboarding: ").font.bold = True
    p.add_run("User connects WhatsApp Business Account via Meta Embedded Signup at /onboarding using the official Meta SDK popup.\n")
    p.add_run("3. Wallet Top-Up: ").font.bold = True
    p.add_run("User adds funds to their prepaid wallet at /billing via Razorpay gateway to cover Meta conversation charges.\n")
    p.add_run("4. Template & Campaign Execution: ").font.bold = True
    p.add_run("User submits template at /templates, gets Meta approval, imports contacts at /contacts, and launches broadcast at /broadcast.")

    add_h2("User Navbar Breakdown & Work Functions")
    p = doc.add_paragraph()
    p.add_run("Below is the complete work breakdown for each navbar item and section available in the User interface:")

    user_nav_data = [
        ("🏠 Dashboard", [
            ("Overview", "/dashboard", "Primary user dashboard displaying live operational statistics: total contacts, messages sent today, unread inbox messages, wallet balance, connected WABA phone number details, and quick launch actions."),
            ("Inbox", "/inbox", "Real-time 2-way WhatsApp Chat Console. Supports viewing active conversations, sending manual text and media replies, selecting canned quick replies, assigning chats to agents, viewing customer profile details, and triggering automated flow bots. Includes unread message badge count.")
        ]),
        ("👥 Contacts", [
            ("All Contacts", "/contacts", "Contact management portal. Enables manual contact creation, bulk CSV file import (mapping name, phone, custom fields), contact search, editing, and single/bulk deletion."),
            ("Labels & Groups", "/contacts", "Organize contacts using color-coded tags and labels (e.g., VIP, Hot Lead, Customer) for organized audience management."),
            ("Segments", "/contacts", "Filter and group contacts into dynamic segments based on tags and activity for targeted marketing.")
        ]),
        ("📢 Campaigns", [
            ("Broadcast", "/broadcast", "Mass WhatsApp broadcast campaign manager. Select target contact labels, pick approved Meta message template, map dynamic body parameters (e.g. {{1}} = Name), attach media header, schedule launch time, and trigger bulk sending."),
            ("Campaign Analytics", "/broadcast", "Track real-time campaign performance metrics: total sent, delivered percentage, read rate, failed count, and delivery error logs."),
            ("Drip Sequences", "/drip-sequences", "Automate multi-day drip campaign sequences (e.g., Day 0 Welcome Message -> Wait 48 Hours -> Day 2 Follow-Up).")
        ]),
        ("📝 Templates", [
            ("My Templates", "/templates", "Manage WhatsApp message templates. View template category (Marketing, Utility, Authentication), language, header format, body content, and live approval status from Meta."),
            ("Pending / Rejected", "/templates", "Track templates currently undergoing Meta review or rejected by Meta with feedback."),
            ("Approved", "/templates", "View ready-to-use Meta-approved templates. Includes interactive Modal to create new text, image, document, or video templates with Call-to-Action buttons and submit directly to Meta Graph API.")
        ]),
        ("🤖 Automation", [
            ("Flow Builder", "/flows", "Drag-and-drop interactive chatbot flow builder. Build multi-branch visual bot flows with message nodes, question prompts, conditional logic, external API webhooks, and Groq AI agent nodes."),
            ("Quick Replies", "/quick-replies", "Create and manage shortcut canned responses (/thanks, /pricing, /address) for rapid agent messaging in live inbox."),
            ("Auto Reply", "/flows", "Configure instant keyword-triggered auto-responses for incoming customer messages."),
            ("Scheduled Messages", "/drip-sequences", "View and schedule single or recurring automated messages to specific contacts.")
        ]),
        ("📈 Analytics", [
            ("Analytics", "/analytics", "Visual charts and graphs showing message volume trends (daily/weekly/monthly), sent vs. received rates, and active contact statistics."),
            ("Campaign ROI", "/analytics", "Track campaign conversion performance and customer response rates."),
            ("Message Stats", "/analytics", "Delivery status breakdown (Delivered %, Read %, Failed %) and failure reason distribution.")
        ]),
        ("💳 Billing & Wallet", [
            ("Wallet & Balance", "/billing", "View current prepaid wallet balance, top up funds securely via Razorpay, view estimated remaining message count based on Meta conversation rates."),
            ("Meta Billing", "/billing", "Track conversation charges by Meta category (Marketing, Utility, Service, Authentication)."),
            ("Payment History", "/billing", "View and download tax invoices for all wallet top-ups and plan payments."),
            ("Subscription Plans", "/plans", "View, upgrade, or downgrade subscription tiers (Starter, Growth, Pro, Enterprise).")
        ]),
        ("📞 WhatsApp", [
            ("Connected Number", "/settings", "View connected WhatsApp phone number, display name, WABA ID, Phone Number ID, and connection status."),
            ("Quality Rating", "/settings", "Monitor Meta phone number quality score (GREEN, YELLOW, RED) and daily messaging limit tier."),
            ("Webhook Status", "/settings", "Verify live webhook synchronization with Meta Cloud API."),
            ("Token Status", "/settings", "Inspect OAuth access token status and trigger manual token refresh if required.")
        ]),
        ("📘 Facebook", [
            ("FB Connection", "/settings", "Trigger Meta Embedded Signup popup modal to link Facebook Business Account and WhatsApp Business Account."),
            ("Business Manager", "/settings", "Check Meta Business Manager verification status."),
            ("Permissions", "/settings", "View granted Meta API scope permissions."),
            ("Token Status", "/settings", "Verify Facebook User Token status.")
        ]),
        ("👨‍💼 Team & Agents", [
            ("Team Members", "/agents", "Invite team members and agents via email to handle customer chat inbox."),
            ("Roles", "/agents", "Assign granular agent roles and inbox permissions."),
            ("Performance", "/agents", "Track agent response speeds, total chats handled, and resolution rates.")
        ]),
        ("⚙️ Settings", [
            ("General Settings", "/settings", "Configure business profile name, default country code, timezone, and language preferences."),
            ("Security & 2FA", "/settings", "Manage password updates and account security settings."),
            ("Notifications", "/settings", "Configure email and WhatsApp alert preferences for low wallet balance and system events."),
            ("Profile & Branding", "/settings", "Upload business logo, update WhatsApp Business profile description, address, email, and website link.")
        ])
    ]

    for section_title, items in user_nav_data:
        add_h3(section_title)
        tbl = doc.add_table(rows=len(items) + 1, cols=3)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header
        h_cols = ["Navbar Item Name", "Route Path", "Functionality & Work Done"]
        for idx, htext in enumerate(h_cols):
            c = tbl.cell(0, idx)
            set_cell_background(c, "0F172A")
            set_cell_margins(c, 100, 100, 120, 120)
            r = c.paragraphs[0].add_run(htext)
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(9.5)

        for row_i, (item_name, route, desc) in enumerate(items, start=1):
            c0 = tbl.cell(row_i, 0)
            c1 = tbl.cell(row_i, 1)
            c2 = tbl.cell(row_i, 2)
            set_cell_margins(c0, 80, 80, 100, 100)
            set_cell_margins(c1, 80, 80, 100, 100)
            set_cell_margins(c2, 80, 80, 100, 100)
            if row_i % 2 == 1:
                set_cell_background(c0, LIGHT_BG)
                set_cell_background(c1, LIGHT_BG)
                set_cell_background(c2, LIGHT_BG)

            r0 = c0.paragraphs[0].add_run(item_name)
            r0.font.bold = True
            r0.font.size = Pt(9.5)
            
            r1 = c1.paragraphs[0].add_run(route)
            r1.font.size = Pt(9)
            r1.font.color.rgb = RED
            
            r2 = c2.paragraphs[0].add_run(desc)
            r2.font.size = Pt(9)

        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # -------------------------------------------------------------
    # SECTION 4: REQUIRED API KEYS & ENVIRONMENT VARIABLES
    # -------------------------------------------------------------
    add_h1("4. Required API Keys & Environment Variables")
    p = doc.add_paragraph()
    p.add_run("Below is the master specification table of all environment variables, API keys, database configurations, and secret tokens required by V ONE CRM across the Server (server/.env) and Client (client/.env):")

    env_table = doc.add_table(rows=22, cols=4)
    env_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ["Variable Name", "Scope", "Description & Purpose", "Current / Example Value"]
    for idx, htext in enumerate(headers):
        c = env_table.cell(0, idx)
        set_cell_background(c, "0F172A")
        set_cell_margins(c, 100, 100, 120, 120)
        r = c.paragraphs[0].add_run(htext)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9.5)

    env_vars_data = [
        ("DB_HOST", "Server", "Hostname for MySQL database connection", "localhost"),
        ("DB_PORT", "Server", "Port number for MySQL database server", "3307 / 3306"),
        ("DB_NAME", "Server", "Database schema name for the application", "whatsapp_crm"),
        ("DB_USER", "Server", "MySQL database user login", "root"),
        ("DB_PASS", "Server", "MySQL database user password", "[empty string or password]"),
        ("JWT_SECRET", "Server", "Secret key for signing & verifying JWT auth tokens", "whatsapp_crm_jwt_secret_2026_k8x9m2"),
        ("WHATSAPP_APP_ID", "Server / Client", "Meta App ID from Facebook Developer Portal", "1590795935988169"),
        ("WHATSAPP_APP_SECRET", "Server", "Meta App Secret for OAuth token exchanges", "f94e6ead41fa1227163240e0f3825ad5"),
        ("WHATSAPP_SYSTEM_USER_TOKEN", "Server", "Permanent System User Access Token for Meta Graph API calls", "EAAWm0gqsuckBSCR..."),
        ("WHATSAPP_WEBHOOK_VERIFY_TOKEN", "Server", "Secret token for Meta Webhook verification handshake", "mahi_crm_webhook_token_2026"),
        ("WHATSAPP_CONFIG_ID", "Server", "Configuration ID for Meta Embedded Signup flow", "1569573811314694"),
        ("WHATSAPP_GRAPH_API_VERSION", "Server", "Meta Graph API version target", "v25.0"),
        ("WHATSAPP_REDIRECT_URI", "Server", "OAuth callback URL for Facebook login", "http://localhost:3000/onboarding/callback"),
        ("WHATSAPP_PHONE_NUMBER_ID", "Server", "Default / Production WhatsApp Phone Number ID", "1269197539606780"),
        ("WHATSAPP_WABA_ID", "Server", "Default / Production WhatsApp Business Account ID", "1014658487838546"),
        ("RAZORPAY_KEY_ID", "Server", "Razorpay Key ID for payments & checkout modal", "rzp_live_T0hHQkCXWwsjRp"),
        ("RAZORPAY_KEY_SECRET", "Server", "Razorpay Secret Key for order signature verification", "nROqQ4okM5617sQ4RrwnKp8v"),
        ("SMTP_HOST", "Server", "Mail server hostname for sending transactional emails", "smtp.gmail.com"),
        ("SMTP_PORT", "Server", "Port number for SMTP mail server", "587"),
        ("SMTP_USERNAME", "Server", "Email address used for SMTP authentication", "kornepatimahankali35@gmail.com"),
        ("SMTP_APP_PASSWORD", "Server", "App-specific password for SMTP server authentication", "kttq onun yugn hwlt"),
        ("SMTP_FROM_NAME", "Server", "Display name for outgoing emails", "V ONE DIGITALS"),
        ("GROQ_API_KEY", "Server", "API key for Groq AI (Llama-3 model for AI Agent bots)", "gsk_ogfbDXmsbmVKX6Uo3L..."),
        ("PORT", "Server", "Port on which Express REST API backend server runs", "5000 / 8000"),
        ("NODE_ENV", "Server", "Node environment mode (development or production)", "development"),
        ("REACT_APP_WHATSAPP_APP_ID", "Client", "Meta App ID exposed to React frontend for Meta SDK", "4257112177765455"),
        ("REACT_APP_API_URL", "Client", "Base REST API URL for React frontend HTTP requests", "http://localhost:8000/api")
    ]

    for row_i, (vname, vscope, vdesc, vexample) in enumerate(env_vars_data, start=1):
        # We need to add rows dynamically if env_table initial size is smaller
        if row_i >= len(env_table.rows):
            env_table.add_row()

        c0 = env_table.cell(row_i, 0)
        c1 = env_table.cell(row_i, 1)
        c2 = env_table.cell(row_i, 2)
        c3 = env_table.cell(row_i, 3)

        set_cell_margins(c0, 80, 80, 100, 100)
        set_cell_margins(c1, 80, 80, 100, 100)
        set_cell_margins(c2, 80, 80, 100, 100)
        set_cell_margins(c3, 80, 80, 100, 100)

        if row_i % 2 == 1:
            set_cell_background(c0, LIGHT_BG)
            set_cell_background(c1, LIGHT_BG)
            set_cell_background(c2, LIGHT_BG)
            set_cell_background(c3, LIGHT_BG)

        r0 = c0.paragraphs[0].add_run(vname)
        r0.font.bold = True
        r0.font.size = Pt(9)
        r0.font.color.rgb = NAVY

        r1 = c1.paragraphs[0].add_run(vscope)
        r1.font.size = Pt(9)
        r1.font.bold = True
        r1.font.color.rgb = RED if vscope == 'Client' else NAVY

        r2 = c2.paragraphs[0].add_run(vdesc)
        r2.font.size = Pt(8.5)

        r3 = c3.paragraphs[0].add_run(vexample)
        r3.font.size = Pt(8.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # Save document
    output_path = os.path.join(r"c:\Users\ramya\Downloads\mahi", "V_ONE_CRM_System_Documentation.docx")
    doc.save(output_path)
    print(f"Successfully created Word document at: {output_path}")

if __name__ == "__main__":
    create_document()
